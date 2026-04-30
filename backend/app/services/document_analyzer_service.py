"""AI-powered document analyzer (V3 fase 7 + V4 fase 1 OCR end-to-end).

Pipeline:
1. Recibimos `bytes` + `content_type` + `filename` + `tipo`.
2. Extraemos texto plano:
   - PDF: `pypdf` primero. Si devuelve <50 chars (señal de PDF imagen),
     fallback a OCR vía `pdf2image` + `pytesseract` (V4 fase 1).
   - DOCX: `python-docx` (soft-fail si no está).
   - imagen (jpeg/png/etc): `pytesseract` directo sobre la imagen.
   - txt/md/csv: decode UTF-8 con fallback latin-1.
3. Truncamos a 6 000 chars antes de mandar al LLM (cost-conscious; ~1.5K
   tokens). Para casos largos, los primeros 6K cubren los headers/datos
   estructurales que es donde viven los campos.
4. Construimos prompt según `tipo`:
   - Schema JSON con los campos esperados.
   - Reglas duras: español, devolver SOLO JSON, null para faltantes.
5. Llamamos a Claude (Messages API, no streaming — esto es one-shot).
   max_tokens=1000.
6. Parseamos: extraemos el primer bloque JSON válido (Claude a veces
   pre-comenta antes del JSON aunque le pidamos lo contrario).
7. Devolvemos `DocumentExtraction` con `extraction_method` y `ocr_pages`.

Privacy: el texto extraído NO se persiste en DB ni en structlog (solo
loggeamos longitudes y tipos, nunca contenido).

Cost: ~$0.003-0.005 por análisis (sonnet 3.5, ~3K tokens input + 500 out).
OCR adicional: 5-10s por página, gratis (binario tesseract local).

V4 fase 1 — OCR fallback:
- DPI 200 (tradeoff calidad/velocidad).
- Lang 'spa+eng' (cubre la mayoría de docs chilenos + nombres en inglés).
- Cap 10 páginas defensivo (un PDF de 100 páginas escaneado tomaría
  varios minutos y bloquearía el endpoint).
- Soft-fail: si tesseract no está instalado a nivel sistema o la lib
  pytesseract no se puede importar, devolvemos extraction_method='failed'
  con warning. NUNCA crasheamos con 500.
"""
from __future__ import annotations

import io
import json
import re
import time
from dataclasses import dataclass
from typing import Any

import structlog

from app.core.config import settings
from app.schemas.document_extraction import DocumentExtraction

log = structlog.get_logger(__name__)

# Tope de chars enviados al LLM. Subir esto multiplica el costo.
MAX_TEXT_CHARS = 6000
MAX_RESPONSE_TOKENS = 1000
MIN_TEXT_CHARS = 20  # bajo esto, el doc es ilegible (PDF de imágenes sin OCR, etc.)

# V4 fase 1 — OCR tunables. Si pypdf devuelve menos chars que este umbral,
# asumimos que el PDF es una imagen escaneada (la mayoría de scans dan al
# menos algunos chars de metadata, pero <50 ya es señal clara).
OCR_MIN_PYPDF_CHARS = 50
# Cap de páginas para OCR (defensivo: 10 páginas a 200 DPI ya tarda 50-100s).
OCR_MAX_PAGES = 10
# DPI: tradeoff. 200 = sweet spot accuracy/speed. 300 mejora calidad ~10% a 2x costo.
OCR_DPI = 200
# Idiomas: spa cubre Chile; eng captura nombres propios y términos legales en inglés.
OCR_LANG = "spa+eng"


class DocumentAnalyzerNotConfigured(Exception):  # noqa: N818
    """`ANTHROPIC_API_KEY` ausente — endpoint debe devolver 503."""


# ---------------------------------------------------------------------------
# Schemas por tipo (el LLM ve estas keys y las llena)
# ---------------------------------------------------------------------------


SCHEMAS: dict[str, dict[str, str]] = {
    "contrato": {
        "contraparte": "Nombre completo de la otra parte (empresa o persona).",
        "rut_contraparte": "RUT de la contraparte si está presente (formato 12.345.678-9).",
        "fecha_inicio": "Fecha de inicio de vigencia (YYYY-MM-DD).",
        "fecha_fin": "Fecha de término o vencimiento (YYYY-MM-DD), null si indefinido.",
        "monto": "Monto total del contrato como número (sin separadores de miles).",
        "moneda": "Código ISO de la moneda (CLP, USD, UF, EUR).",
        "descripcion": "Resumen 1-2 frases del objeto del contrato.",
        "partes": "Lista de partes intervinientes (string corto, separadas por coma).",
    },
    "f29": {
        "empresa": "Razón social del contribuyente.",
        "rut_empresa": "RUT del contribuyente (formato 76.123.456-7).",
        "periodo_tributario": "Período en formato MM_AA (ej: 02_26 para feb 2026).",
        "fecha_vencimiento": "Fecha de vencimiento del pago (YYYY-MM-DD).",
        "monto_a_pagar": "Monto total a pagar en CLP, número entero sin separadores.",
        "estado": "Uno de: pendiente, pagado, vencido, exento.",
    },
    "trabajador_contrato": {
        "nombre_completo": "Nombre completo del trabajador.",
        "rut": "RUT del trabajador (12345678-9).",
        "cargo": "Cargo o título del puesto.",
        "fecha_ingreso": "Fecha de inicio de la relación laboral (YYYY-MM-DD).",
        "sueldo_bruto": "Sueldo bruto mensual en CLP, número entero.",
        "tipo_contrato": "Uno de: indefinido, plazo_fijo, honorarios, part_time.",
        "email": "Email del trabajador si está presente.",
        "telefono": "Teléfono del trabajador si está presente.",
    },
    "factura": {
        "proveedor_rut": "RUT del proveedor/emisor.",
        "proveedor_nombre": "Razón social del proveedor.",
        "numero_factura": "Folio o número de factura.",
        "fecha": "Fecha de emisión (YYYY-MM-DD).",
        "monto_neto": "Monto neto sin IVA, número entero.",
        "iva": "Monto del IVA, número entero.",
        "total": "Total con IVA incluido, número entero.",
        "descripcion": "Glosa o detalle resumido del bien/servicio.",
    },
    "liquidacion": {
        "trabajador_nombre": "Nombre completo del trabajador.",
        "rut": "RUT del trabajador.",
        "periodo": "Período liquidado en formato MM_AA o 'YYYY-MM'.",
        "sueldo_bruto": "Sueldo bruto del período en CLP.",
        "descuentos": "Total de descuentos legales y otros en CLP.",
        "liquido_pagar": "Líquido a pagar en CLP.",
    },
}


# ---------------------------------------------------------------------------
# Text extraction
# ---------------------------------------------------------------------------


@dataclass(frozen=True, slots=True)
class _ExtractResult:
    text: str
    warnings: tuple[str, ...] = ()
    # V4 fase 1: cómo se sacó el texto (pypdf | ocr | hybrid | docx | text |
    # image_ocr | failed | unknown). El endpoint pasa esto al schema response.
    method: str = "unknown"
    # Páginas que pasaron por OCR (None si no aplicó).
    ocr_pages: int | None = None


def _decode_text(content: bytes) -> str:
    try:
        return content.decode("utf-8")
    except UnicodeDecodeError:
        return content.decode("latin-1", errors="ignore")


def extract_text_pdf(content: bytes) -> str:
    """Extrae texto de un PDF con pypdf, tolerante a páginas corruptas.

    Para PDFs escaneados (sin texto extraíble) prefiera
    `extract_text_pdf_with_fallback`, que intenta OCR si pypdf no rinde.
    """
    from pypdf import PdfReader

    reader = PdfReader(io.BytesIO(content))
    parts: list[str] = []
    for page in reader.pages:
        try:
            parts.append(page.extract_text() or "")
        except Exception as exc:
            log.warning("doc_analyzer.pdf_page_failed", error=str(exc))
    out = "\n\n".join(p for p in parts if p.strip())
    return re.sub(r"\n{3,}", "\n\n", out)


def _ocr_pdf_pages(content: bytes) -> tuple[str, int, list[str]]:
    """OCR de un PDF página-por-página vía pdf2image + pytesseract.

    Devuelve `(texto_concatenado, pages_ocr, warnings)`. Soft-fail:
    si las libs/binarios faltan, devuelve `("", 0, [warning])` — nunca raise.
    """
    warnings: list[str] = []
    try:
        import pytesseract  # type: ignore[import-not-found]
        from pdf2image import convert_from_bytes  # type: ignore[import-not-found]
    except ImportError as exc:
        warnings.append(
            "OCR no disponible: instalá pytesseract + pdf2image + el binario "
            f"tesseract y poppler-utils a nivel sistema ({exc})."
        )
        log.warning("doc_analyzer.ocr_import_failed", error=str(exc))
        return "", 0, warnings

    try:
        images = convert_from_bytes(
            content,
            dpi=OCR_DPI,
            first_page=1,
            last_page=OCR_MAX_PAGES,
        )
    except Exception as exc:
        # pdf2image lanza PDFInfoNotInstalledError, PDFPageCountError, etc.
        # cuando poppler falta o el PDF está corrupto. Soft-fail.
        warnings.append(
            f"No pude convertir el PDF a imágenes para OCR (¿poppler-utils "
            f"instalado?): {exc}"
        )
        log.warning("doc_analyzer.pdf2image_failed", error=str(exc))
        return "", 0, warnings

    parts: list[str] = []
    pages_ocr = 0
    for idx, img in enumerate(images, start=1):
        try:
            page_text = pytesseract.image_to_string(img, lang=OCR_LANG)
        except Exception as exc:
            warnings.append(f"OCR falló en página {idx}: {exc}")
            log.warning("doc_analyzer.ocr_page_failed", page=idx, error=str(exc))
            continue
        if page_text and page_text.strip():
            parts.append(page_text)
        pages_ocr += 1

    text_out = "\n\n".join(parts)
    return re.sub(r"\n{3,}", "\n\n", text_out), pages_ocr, warnings


def extract_text_pdf_with_fallback(content: bytes) -> tuple[str, dict[str, Any]]:
    """Extrae texto de PDF con fallback automático a OCR si pypdf no rinde.

    Estrategia (V4 fase 1):
    1. Intenta pypdf. Si devuelve >= OCR_MIN_PYPDF_CHARS, listo (PDF digital).
    2. Si <50 chars (PDF imagen / escaneo): convierte páginas a imágenes y
       pasa cada una por tesseract (`spa+eng`, DPI 200, cap 10 páginas).
    3. Si pypdf devolvió algo y OCR también: marca como 'hybrid' y queda
       con el OCR (que típicamente es más completo).
    4. Soft-fail: si tesseract/poppler no están, devuelve lo poco que dio
       pypdf con method='failed' y un warning.

    Returns:
        (text, meta) donde meta tiene keys:
          - method: 'pypdf' | 'ocr' | 'hybrid' | 'failed'
          - pages_ocr: int (cantidad de páginas que pasaron por OCR; 0 si no)
          - pages_total: int (cantidad total de páginas en el PDF)
          - ocr_duration_ms: float (latencia del paso OCR; 0 si no aplicó)
          - warnings: list[str]
    """
    warnings: list[str] = []
    pypdf_text = ""
    pages_total = 0

    try:
        from pypdf import PdfReader

        reader = PdfReader(io.BytesIO(content))
        pages_total = len(reader.pages)
        pypdf_text = extract_text_pdf(content)
    except Exception as exc:
        warnings.append(f"pypdf no pudo leer el PDF: {exc}")
        log.warning("doc_analyzer.pypdf_failed", error=str(exc))

    # Caso happy path: PDF digital con texto suficiente.
    if len(pypdf_text.strip()) >= OCR_MIN_PYPDF_CHARS:
        log.info(
            "doc_analyzer.pdf.pypdf_ok",
            chars=len(pypdf_text),
            pages_total=pages_total,
        )
        return pypdf_text, {
            "method": "pypdf",
            "pages_ocr": 0,
            "pages_total": pages_total,
            "ocr_duration_ms": 0.0,
            "warnings": warnings,
        }

    # Fallback OCR. Loggeamos el motivo para observabilidad de costo.
    log.info(
        "doc_analyzer.pdf.ocr_fallback",
        pypdf_chars=len(pypdf_text.strip()),
        threshold=OCR_MIN_PYPDF_CHARS,
        pages_total=pages_total,
    )
    t0 = time.perf_counter()
    ocr_text, pages_ocr, ocr_warnings = _ocr_pdf_pages(content)
    duration_ms = (time.perf_counter() - t0) * 1000.0

    log.info(
        "doc_analyzer.pdf.ocr_done",
        ocr_duration_ms=round(duration_ms, 1),
        pages_ocr=pages_ocr,
        pages_total=pages_total,
        ocr_chars=len(ocr_text),
    )

    warnings.extend(ocr_warnings)

    if not ocr_text.strip():
        # OCR no rindió (tesseract ausente, PDF dañado, etc.) — devolvemos
        # lo poquito que pypdf sacó (o vacío) con method='failed'.
        return pypdf_text, {
            "method": "failed",
            "pages_ocr": 0,
            "pages_total": pages_total,
            "ocr_duration_ms": round(duration_ms, 1),
            "warnings": warnings,
        }

    # Si pypdf también rindió algo (caso raro: PDFs mixtos con páginas
    # digitales + escaneadas), marcamos 'hybrid' y devolvemos el OCR (más
    # completo). Si pypdf no dio nada, es OCR puro.
    method = "hybrid" if pypdf_text.strip() else "ocr"
    return ocr_text, {
        "method": method,
        "pages_ocr": pages_ocr,
        "pages_total": pages_total,
        "ocr_duration_ms": round(duration_ms, 1),
        "warnings": warnings,
    }


def extract_text_docx(content: bytes) -> str:
    """Extrae texto de un .docx usando python-docx."""
    try:
        from docx import Document  # type: ignore[import-not-found]
    except ImportError as exc:  # pragma: no cover
        raise RuntimeError(
            "python-docx no instalado — agregar a deps de backend"
        ) from exc

    doc = Document(io.BytesIO(content))
    parts = [p.text for p in doc.paragraphs if p.text and p.text.strip()]
    # Tablas (común en contratos laborales con grilla de remuneración).
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text and c.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    return "\n".join(parts)


def extract_text_image(
    content: bytes,
    mime: str | None = None,
) -> tuple[str | None, list[str]]:
    """OCR directo sobre una imagen (jpeg/png/tiff/etc.) con pytesseract.

    Soft-fail si la lib o el binario tesseract no están instalados —
    devuelve `(None, [warning])` en lugar de crashear.

    El parámetro `mime` es informativo (usado para logging); Pillow detecta
    el formato real desde los magic bytes.
    """
    warnings: list[str] = []
    try:
        import pytesseract  # type: ignore[import-not-found]
        from PIL import Image  # type: ignore[import-not-found]
    except ImportError as exc:
        warnings.append(
            "OCR no disponible: instalá pytesseract + Pillow + el binario "
            f"tesseract a nivel sistema ({exc})."
        )
        log.warning("doc_analyzer.ocr_image_import_failed", error=str(exc))
        return None, warnings

    t0 = time.perf_counter()
    try:
        img = Image.open(io.BytesIO(content))
        text_extracted = pytesseract.image_to_string(img, lang=OCR_LANG)
    except Exception as exc:
        log.warning("doc_analyzer.ocr_image_failed", error=str(exc), mime=mime)
        warnings.append(f"OCR falló: {exc}")
        return None, warnings

    duration_ms = (time.perf_counter() - t0) * 1000.0
    log.info(
        "doc_analyzer.image_ocr.done",
        ocr_duration_ms=round(duration_ms, 1),
        ocr_chars=len(text_extracted) if text_extracted else 0,
        mime=mime,
    )

    if not text_extracted or not text_extracted.strip():
        warnings.append("OCR devolvió texto vacío.")
        return None, warnings
    return text_extracted, warnings


def _ext_from_content_type(content_type: str, filename: str | None) -> str:
    """Devuelve un sufijo lowercase tipo 'pdf'/'docx'/'png' según content-type+nombre."""
    ct = (content_type or "").lower()
    if "pdf" in ct:
        return "pdf"
    if "officedocument.wordprocessingml" in ct or ct.endswith("/msword"):
        return "docx"
    if "png" in ct:
        return "png"
    if "jpeg" in ct or "jpg" in ct:
        return "jpg"
    if "plain" in ct or "markdown" in ct or "csv" in ct:
        return "txt"
    # Fallback al filename
    if filename and "." in filename:
        return filename.rsplit(".", 1)[-1].lower()
    return ""


async def extract_text(
    content: bytes,
    content_type: str,
    filename: str | None = None,
) -> _ExtractResult:
    """Dispatcher: elige extractor según content-type/extensión.

    V4 fase 1:
    - PDF usa `extract_text_pdf_with_fallback` (pypdf → OCR si <50 chars).
    - Imágenes van directo a OCR.
    - El `_ExtractResult` incluye `method` y `ocr_pages` para que el
      endpoint los exponga en la respuesta y el frontend muestre el chip
      "OCR aplicado" cuando corresponda.
    """
    ext = _ext_from_content_type(content_type, filename)
    if ext == "pdf":
        text, meta = extract_text_pdf_with_fallback(content)
        return _ExtractResult(
            text=text,
            warnings=tuple(meta.get("warnings") or []),
            method=str(meta.get("method", "pypdf")),
            ocr_pages=meta.get("pages_ocr") or None,
        )
    if ext == "docx":
        return _ExtractResult(text=extract_text_docx(content), method="docx")
    if ext in {"png", "jpg", "jpeg", "gif", "webp", "tif", "tiff"}:
        text_out, warns = extract_text_image(content, mime=content_type)
        # Si OCR rindió → 'image_ocr'; si soft-failed → 'failed'.
        method = "image_ocr" if text_out else "failed"
        return _ExtractResult(
            text=text_out or "",
            warnings=tuple(warns),
            method=method,
            ocr_pages=1 if text_out else None,
        )
    if ext in {"txt", "md", "csv", "log", ""}:
        return _ExtractResult(text=_decode_text(content), method="text")
    # Tipo desconocido — intenta como texto best-effort.
    return _ExtractResult(
        text=_decode_text(content),
        warnings=(f"Tipo de archivo no reconocido ({ext or content_type}); intenté como texto.",),
        method="text",
    )


# ---------------------------------------------------------------------------
# Prompt building
# ---------------------------------------------------------------------------


def _schema_block(tipo: str) -> str:
    """Render del schema como JSON-prima legible para el LLM."""
    if tipo == "auto":
        # Listamos todos los tipos para que el LLM elija.
        lines = ["Tipos posibles:"]
        for key, schema in SCHEMAS.items():
            lines.append(f"\n- {key}:")
            for k, desc in schema.items():
                lines.append(f"    {k}: {desc}")
        return "\n".join(lines)
    schema = SCHEMAS.get(tipo, {})
    parts = ["{"]
    for k, desc in schema.items():
        parts.append(f'  "{k}": <{desc}>,')
    parts.append("}")
    return "\n".join(parts)


def build_prompt(text: str, tipo: str) -> str:
    """Prompt user-side. Se acompaña con un system prompt estricto."""
    truncated = text[:MAX_TEXT_CHARS]
    truncation_note = (
        f"\n\n[Texto truncado a {MAX_TEXT_CHARS} caracteres. Total original: "
        f"{len(text)} chars.]"
        if len(text) > MAX_TEXT_CHARS
        else ""
    )

    if tipo == "auto":
        return (
            "Analizá el siguiente documento. Detectá su tipo (contrato, f29, "
            "trabajador_contrato, factura, liquidacion) y extraé los campos "
            "estructurados correspondientes.\n\n"
            "DOCUMENTO:\n"
            f"{truncated}{truncation_note}\n\n"
            "Devolvé EXCLUSIVAMENTE un JSON con esta forma exacta:\n"
            "{\n"
            '  "tipo_detectado": "<uno de los tipos>",\n'
            '  "confidence": <float 0-1>,\n'
            '  "fields": { <campos del schema correspondiente, null si faltan> },\n'
            '  "warnings": [ "<nota si hay ambigüedades>" ]\n'
            "}\n\n"
            f"Schemas por tipo:\n{_schema_block('auto')}\n\n"
            "JSON:"
        )

    return (
        f"Analizá este documento (tipo: {tipo}) y extraé los campos estructurados.\n\n"
        "DOCUMENTO:\n"
        f"{truncated}{truncation_note}\n\n"
        "Devolvé EXCLUSIVAMENTE un JSON con esta forma exacta:\n"
        "{\n"
        f'  "tipo_detectado": "{tipo}",\n'
        '  "confidence": <float 0-1>,\n'
        '  "fields": <objeto con los campos del schema; null para los ausentes>,\n'
        '  "warnings": [ "<nota si hay ambigüedades>" ]\n'
        "}\n\n"
        f"Schema esperado para `fields`:\n{_schema_block(tipo)}\n\n"
        "Reglas:\n"
        "- Si un campo no está en el documento, devolvé `null` (no inventes).\n"
        "- Fechas en formato YYYY-MM-DD.\n"
        "- Montos como números enteros sin separadores ni símbolo de moneda.\n"
        "- Si dudás del valor, agregá una nota a `warnings`.\n\n"
        "JSON:"
    )


SYSTEM_PROMPT = (
    "Sos un extractor de datos estructurados de documentos financieros y "
    "legales chilenos para Cehta Capital. Respondés SIEMPRE con un JSON "
    "válido, sin texto antes ni después, sin markdown fences. No inventes "
    "datos: si un campo no está, usá null. Las fechas van en formato "
    "YYYY-MM-DD; los RUT en formato chileno (12.345.678-9 o 12345678-9). "
    "Si el documento está en otro idioma, traducí solo los valores estructurados "
    "(no los descriptivos) al español neutro."
)


# ---------------------------------------------------------------------------
# Claude call + JSON parsing
# ---------------------------------------------------------------------------


def _anthropic_client() -> Any:
    if not settings.anthropic_api_key:
        raise DocumentAnalyzerNotConfigured(
            "ANTHROPIC_API_KEY no configurado — el AI Document Analyzer "
            "no está disponible. Configurá la key o subí el archivo "
            "manualmente sin auto-fill."
        )
    from anthropic import AsyncAnthropic

    return AsyncAnthropic(api_key=settings.anthropic_api_key)


_JSON_OBJECT_RE = re.compile(r"\{.*\}", re.DOTALL)


def parse_llm_json(raw: str) -> dict[str, Any]:
    """Parser tolerante: extrae el primer bloque JSON {…} del texto del LLM.

    Por qué tolerante: Claude a veces antepone "Aquí está el JSON:" aunque le
    pidamos lo contrario, o envuelve en ```json ... ```. Esta función:
    1. Limpia fences markdown (```json … ```).
    2. Busca el primer `{` … `}` balanceado.
    3. Hace `json.loads`. Si falla, raise.
    """
    cleaned = raw.strip()
    # Remove ```json fences
    cleaned = re.sub(r"^```(?:json)?\s*", "", cleaned)
    cleaned = re.sub(r"\s*```$", "", cleaned)

    # Try direct parse first
    try:
        return json.loads(cleaned)  # type: ignore[no-any-return]
    except json.JSONDecodeError:
        pass

    # Fallback: balanced-brace extraction
    start = cleaned.find("{")
    if start == -1:
        raise ValueError("No JSON object found in LLM response")
    depth = 0
    end = -1
    for i in range(start, len(cleaned)):
        ch = cleaned[i]
        if ch == "{":
            depth += 1
        elif ch == "}":
            depth -= 1
            if depth == 0:
                end = i + 1
                break
    if end == -1:
        raise ValueError("Unbalanced braces in LLM response")
    return json.loads(cleaned[start:end])  # type: ignore[no-any-return]


def _normalize_extraction(
    parsed: dict[str, Any],
    *,
    requested_tipo: str,
    raw_text: str,
    extraction_warnings: list[str],
    extraction_method: str | None = None,
    ocr_pages: int | None = None,
) -> DocumentExtraction:
    """Coerce el dict del LLM al shape de DocumentExtraction, con defaults seguros."""
    tipo_detectado = str(
        parsed.get("tipo_detectado")
        or (requested_tipo if requested_tipo != "auto" else "desconocido")
    )

    # confidence robust: a veces el LLM devuelve "0.9", "90%", o lo omite.
    raw_conf = parsed.get("confidence")
    confidence: float
    if isinstance(raw_conf, int | float):
        confidence = float(raw_conf)
    elif isinstance(raw_conf, str):
        s = raw_conf.strip().rstrip("%")
        try:
            confidence = float(s)
            if confidence > 1.0:
                confidence /= 100.0
        except ValueError:
            confidence = 0.5
    else:
        confidence = 0.5
    confidence = max(0.0, min(1.0, confidence))

    fields = parsed.get("fields")
    if not isinstance(fields, dict):
        fields = {}

    warnings_from_llm = parsed.get("warnings") or []
    if not isinstance(warnings_from_llm, list):
        warnings_from_llm = []
    warnings_combined = [str(w) for w in warnings_from_llm if w] + extraction_warnings

    if confidence < 0.5:
        warnings_combined.append(
            "Análisis con baja confianza — revisá los campos antes de guardar."
        )

    return DocumentExtraction(
        tipo_detectado=tipo_detectado,
        confidence=confidence,
        fields=fields,
        raw_text_preview=raw_text[:500],
        warnings=warnings_combined,
        extraction_method=extraction_method,
        ocr_pages=ocr_pages,
    )


async def analyze_document(
    text: str,
    tipo: str,
    *,
    filename: str | None = None,
    extraction_warnings: list[str] | None = None,
    extraction_method: str | None = None,
    ocr_pages: int | None = None,
) -> DocumentExtraction:
    """Llama a Claude y devuelve un `DocumentExtraction`.

    `extraction_method` y `ocr_pages` los pasa el endpoint desde el
    `_ExtractResult` para que la respuesta lleve traza de cómo se obtuvo
    el texto (frontend muestra chip "OCR aplicado" cuando corresponde).
    """
    extraction_warnings = list(extraction_warnings or [])

    client = _anthropic_client()
    prompt = build_prompt(text, tipo)

    log.info(
        "doc_analyzer.analyze.start",
        tipo=tipo,
        filename=filename,
        text_chars=len(text),
        truncated_to=min(len(text), MAX_TEXT_CHARS),
    )

    try:
        message = await client.messages.create(
            model=settings.ai_chat_model,
            max_tokens=MAX_RESPONSE_TOKENS,
            system=SYSTEM_PROMPT,
            messages=[{"role": "user", "content": prompt}],
        )
    except Exception as exc:
        log.error("doc_analyzer.llm_failed", error=str(exc))
        raise

    # Anthropic SDK: content es lista de blocks; agarramos el primer text block.
    raw_text = ""
    for block in getattr(message, "content", []) or []:
        text_attr = getattr(block, "text", None)
        if text_attr:
            raw_text = text_attr
            break

    if not raw_text:
        raise ValueError("El LLM devolvió respuesta vacía.")

    try:
        parsed = parse_llm_json(raw_text)
    except (ValueError, json.JSONDecodeError) as exc:
        log.warning("doc_analyzer.json_parse_failed", error=str(exc))
        # Devolvemos un fallback "vacío" en lugar de explotar — el frontend
        # muestra warning y deja que el usuario llene a mano.
        return DocumentExtraction(
            tipo_detectado=tipo if tipo != "auto" else "desconocido",
            confidence=0.0,
            fields={},
            raw_text_preview=text[:500],
            warnings=[
                "No pude parsear la respuesta del modelo como JSON. "
                "Llená los campos manualmente.",
                *extraction_warnings,
            ],
            extraction_method=extraction_method,
            ocr_pages=ocr_pages,
        )

    result = _normalize_extraction(
        parsed,
        requested_tipo=tipo,
        raw_text=text,
        extraction_warnings=extraction_warnings,
        extraction_method=extraction_method,
        ocr_pages=ocr_pages,
    )
    log.info(
        "doc_analyzer.analyze.ok",
        tipo_detectado=result.tipo_detectado,
        confidence=result.confidence,
        field_count=len(result.fields),
        warning_count=len(result.warnings),
    )
    return result
